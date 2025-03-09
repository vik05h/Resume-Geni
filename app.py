import os
import sqlite3
import docx
import PyPDF2
from flask import Flask, request, render_template, redirect, url_for, flash
from urllib.parse import quote as url_quote  # FIXED: Replaced werkzeug.urls.url_quote with urllib.parse.quote
from transformers import pipeline

# ------------------------------------------------------------------------------
# 1. Configure Flask App
# ------------------------------------------------------------------------------
app = Flask(__name__)
app.secret_key = "some_secret_key"  # Replace with a secure key in production

# Configure file uploads
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# ------------------------------------------------------------------------------
# 2. SQLite Database Setup
# ------------------------------------------------------------------------------
DB_NAME = 'resume.db'

def init_db():
    """Initialize DB and create table if it doesn't exist."""
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS resumes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT,
            content TEXT,
            score INTEGER,
            suggestions TEXT
        )
    """)
    conn.commit()
    conn.close()

init_db()  # Ensure DB is ready on startup

# ------------------------------------------------------------------------------
# 3. Utility function to extract text from PDF or DOCX
# ------------------------------------------------------------------------------
def extract_text_from_file(filepath):
    file_text = ""

    if filepath.lower().endswith(".pdf"):
        try:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    file_text += page.extract_text() or ""
        except Exception as e:
            print("Error reading PDF:", e)

    elif filepath.lower().endswith(".docx"):
        try:
            doc = docx.Document(filepath)  # Use Document() correctly
            for para in doc.paragraphs:
                file_text += para.text + "\n"
        except Exception as e:
            print("Error reading DOCX:", e)

    else:
        # Handle text files if needed
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            file_text = f.read()

    return file_text.strip()  # ✅ Fix: Ensure a valid string is returned



# ------------------------------------------------------------------------------
# 4. A simple scoring system (heuristic) + suggestions
# ------------------------------------------------------------------------------
def score_resume(text):
    """
    This is a mock scoring function. 
    In reality, you can do NLP or more advanced checks:
      - Check presence of keywords (skills, education, etc.)
      - Grammar checks
      - Format (section headings, bullet points, etc.)
    We'll just do something naive here.
    """
    score = 0
    suggestions = []

    if "experience" in text.lower():
        score += 20
    else:
        suggestions.append("Consider adding a dedicated 'Experience' section.")

    if "education" in text.lower():
        score += 20
    else:
        suggestions.append("Provide an 'Education' section with degrees, institutions, and dates.")

    if "project" in text.lower() or "projects" in text.lower():
        score += 20
    else:
        suggestions.append("List any relevant projects, with brief descriptions.")

    skill_keywords = ["python", "java", "flask", "sql", "aws", "javascript", "react", "css", "html"]
    found_skill = any(kw in text.lower() for kw in skill_keywords)
    if found_skill:
        score += 20
    else:
        suggestions.append("Highlight your technical skills (Python, Java, AWS, etc.).")

    word_count = len(text.split())
    if word_count < 100:
        suggestions.append("Resume seems short, consider adding more relevant details.")
    elif word_count > 1000:
        suggestions.append("Resume might be too long, consider summarizing or condensing info.")
    else:
        score += 20

    if score > 100:
        score = 100

    return score, suggestions

def ai_score_resume(text):
    """
    Uses a pre-trained NLP model to score resumes based on content relevance.
    """
    try:
        summarizer = pipeline("text-classification", model="bert-base-uncased")  # Change this model if needed
        ai_result = summarizer(text)

        # Extract score (normalize to 100)
        score = int(ai_result[0]['score'] * 100)
        suggestions = ["Consider improving the clarity and structure of your resume based on AI analysis."]

        return score, suggestions
    except Exception as e:
        print(f"Error in AI scoring: {e}")
        return 50, ["AI model could not process resume. Falling back to default score."]


# ------------------------------------------------------------------------------
# 5. Flask Routes
# ------------------------------------------------------------------------------
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'resume' not in request.files:
        flash("No file part in the request.")
        return redirect(url_for('index'))

    resume_file = request.files['resume']
    if resume_file.filename == '':
        flash("No selected file.")
        return redirect(url_for('index'))

    # Get the user's selection: normal or AI scoring
    scoring_method = request.form.get("scoring_method", "normal")

    filename = url_quote(resume_file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    resume_file.save(filepath)

    text = extract_text_from_file(filepath)

    if not text:
        flash("Error extracting text from resume. Please upload a valid PDF or DOCX.")
        return redirect(url_for('index'))

    # Select scoring method
    if scoring_method == "normal":
        score, suggestions_list = score_resume(text)  # Hardcoded scoring
    else:
        score, suggestions_list = ai_score_resume(text)  # AI model-based scoring

    suggestions_str = "\n".join(suggestions_list)

    # Save results to DB
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO resumes (filename, content, score, suggestions)
        VALUES (?, ?, ?, ?)
    """, (filename, text, score, suggestions_str))
    conn.commit()
    conn.close()

    return redirect(url_for('result', resume_id=cursor.lastrowid))


@app.route('/result/<int:resume_id>')
def result(resume_id):
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute("SELECT filename, score, suggestions FROM resumes WHERE id = ?", (resume_id,))
    row = cursor.fetchone()
    conn.close()

    if row:
        filename, score, suggestions = row
        suggestions_list = suggestions.split("\n") if suggestions else []
        return render_template('result.html',
                               filename=filename,
                               score=score,
                               suggestions=suggestions_list)
    else:
        return "Resume not found.", 404

# ------------------------------------------------------------------------------
# 6. Run Flask
# ------------------------------------------------------------------------------
if __name__ == '__main__':
    app.run(debug=True)